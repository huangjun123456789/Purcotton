# -*- coding: utf-8 -*-
"""热力图分析报告API"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy import select, func
from app.database import AsyncSessionLocal
from app.models.warehouse import (
    Warehouse, Zone, Aisle, Shelf, Location, LocationHeatData
)
import os
from datetime import datetime
from typing import Optional

router = APIRouter()


async def fetch_report_data(zone_id: Optional[int] = None):
    """获取报告数据，支持按库区筛选"""
    async with AsyncSessionLocal() as db:
        data = {}
        
        # 获取库区信息（用于报告标题）
        if zone_id:
            zone_result = await db.execute(
                select(Zone.name, Warehouse.name)
                .join(Warehouse, Zone.warehouse_id == Warehouse.id)
                .where(Zone.id == zone_id)
            )
            zone_info = zone_result.fetchone()
            if zone_info:
                data['zone_name'] = zone_info[0]
                data['warehouse_name'] = zone_info[1]
                data['report_scope'] = f"{zone_info[1]} - {zone_info[0]}"
            else:
                data['report_scope'] = "全部"
        else:
            data['report_scope'] = "全部"
        
        # 构建库位ID子查询（用于筛选）
        if zone_id:
            # 获取该库区下所有库位的ID
            location_ids_query = (
                select(Location.id)
                .join(Shelf, Location.shelf_id == Shelf.id)
                .join(Aisle, Shelf.aisle_id == Aisle.id)
                .where(Aisle.zone_id == zone_id)
            )
        else:
            location_ids_query = None
        
        # 基础统计 - 根据是否有zone_id来决定统计范围
        if zone_id:
            # 只统计指定库区
            data['warehouse_count'] = 1
            data['zone_count'] = 1
            data['aisle_count'] = (await db.execute(
                select(func.count(Aisle.id)).where(Aisle.zone_id == zone_id)
            )).scalar() or 0
            data['shelf_count'] = (await db.execute(
                select(func.count(Shelf.id))
                .join(Aisle, Shelf.aisle_id == Aisle.id)
                .where(Aisle.zone_id == zone_id)
            )).scalar() or 0
            data['location_count'] = (await db.execute(
                select(func.count(Location.id))
                .join(Shelf, Location.shelf_id == Shelf.id)
                .join(Aisle, Shelf.aisle_id == Aisle.id)
                .where(Aisle.zone_id == zone_id)
            )).scalar() or 0
            # 有热力数据的库位数
            data['active_location_count'] = (await db.execute(
                select(func.count(func.distinct(LocationHeatData.location_id)))
                .where(LocationHeatData.location_id.in_(location_ids_query))
            )).scalar() or 0
            data['heat_data_count'] = (await db.execute(
                select(func.count(LocationHeatData.id))
                .where(LocationHeatData.location_id.in_(location_ids_query))
            )).scalar() or 0
        else:
            # 统计全部
            data['warehouse_count'] = (await db.execute(select(func.count(Warehouse.id)))).scalar() or 0
            data['zone_count'] = (await db.execute(select(func.count(Zone.id)))).scalar() or 0
            data['aisle_count'] = (await db.execute(select(func.count(Aisle.id)))).scalar() or 0
            data['shelf_count'] = (await db.execute(select(func.count(Shelf.id)))).scalar() or 0
            data['location_count'] = (await db.execute(select(func.count(Location.id)))).scalar() or 0
            # 有热力数据的库位数
            data['active_location_count'] = (await db.execute(
                select(func.count(func.distinct(LocationHeatData.location_id)))
            )).scalar() or 0
            data['heat_data_count'] = (await db.execute(select(func.count(LocationHeatData.id)))).scalar() or 0
        
        # 仓库列表
        warehouses_result = await db.execute(
            select(Warehouse.id, Warehouse.code, Warehouse.name)
        )
        data['warehouses'] = [
            {'id': r[0], 'code': r[1], 'name': r[2]} 
            for r in warehouses_result.fetchall()
        ]
        
        # 热度统计 - 筛选指定库区
        if zone_id:
            heat_stats = await db.execute(
                select(
                    func.min(LocationHeatData.heat_value),
                    func.max(LocationHeatData.heat_value),
                    func.avg(LocationHeatData.heat_value)
                ).where(LocationHeatData.location_id.in_(location_ids_query))
            )
        else:
            heat_stats = await db.execute(
                select(
                    func.min(LocationHeatData.heat_value),
                    func.max(LocationHeatData.heat_value),
                    func.avg(LocationHeatData.heat_value)
                )
            )
        row = heat_stats.fetchone()
        data['heat_stats'] = {
            'min': row[0] or 0,
            'max': row[1] or 0,
            'avg': row[2] or 0
        }
        
        # 拣货频率统计 - 筛选指定库区
        if zone_id:
            freq_stats = await db.execute(
                select(
                    func.min(LocationHeatData.pick_frequency),
                    func.max(LocationHeatData.pick_frequency),
                    func.avg(LocationHeatData.pick_frequency)
                ).where(LocationHeatData.location_id.in_(location_ids_query))
            )
        else:
            freq_stats = await db.execute(
                select(
                    func.min(LocationHeatData.pick_frequency),
                    func.max(LocationHeatData.pick_frequency),
                    func.avg(LocationHeatData.pick_frequency)
                )
            )
        row = freq_stats.fetchone()
        data['freq_stats'] = {
            'min': row[0] or 0,
            'max': row[1] or 0,
            'avg': row[2] or 0
        }
        
        # 热度分布 - 筛选指定库区
        if zone_id:
            heat_dist_result = await db.execute(
                select(LocationHeatData.heat_value)
                .where(LocationHeatData.location_id.in_(location_ids_query))
            )
        else:
            heat_dist_result = await db.execute(
                select(LocationHeatData.heat_value)
            )
        heat_values = [r[0] for r in heat_dist_result.fetchall()]
        
        # 手动计算分布
        distribution = {
            '极冷 (0-50)': 0,
            '较冷 (50-100)': 0,
            '一般 (100-200)': 0,
            '较热 (200-300)': 0,
            '极热 (>300)': 0
        }
        for v in heat_values:
            if v < 50:
                distribution['极冷 (0-50)'] += 1
            elif v < 100:
                distribution['较冷 (50-100)'] += 1
            elif v < 200:
                distribution['一般 (100-200)'] += 1
            elif v < 300:
                distribution['较热 (200-300)'] += 1
            else:
                distribution['极热 (>300)'] += 1
        
        data['heat_distribution'] = [
            {'heat_level': k, 'cnt': v} for k, v in distribution.items()
        ]
        
        # TOP热门库位 - 筛选指定库区（按库位分组去重）
        if zone_id:
            top_hot_result = await db.execute(
                select(
                    Location.full_code,
                    func.max(LocationHeatData.heat_value),
                    func.sum(LocationHeatData.pick_frequency),
                    func.avg(LocationHeatData.turnover_rate)
                ).join(Location, LocationHeatData.location_id == Location.id)
                .where(LocationHeatData.location_id.in_(location_ids_query))
                .group_by(Location.id, Location.full_code)
                .order_by(func.max(LocationHeatData.heat_value).desc())
                .limit(20)
            )
        else:
            top_hot_result = await db.execute(
                select(
                    Location.full_code,
                    func.max(LocationHeatData.heat_value),
                    func.sum(LocationHeatData.pick_frequency),
                    func.avg(LocationHeatData.turnover_rate)
                ).join(Location, LocationHeatData.location_id == Location.id)
                .group_by(Location.id, Location.full_code)
                .order_by(func.max(LocationHeatData.heat_value).desc())
                .limit(20)
            )
        data['top_hot'] = [
            {'full_code': r[0], 'heat_value': r[1] or 0, 'pick_frequency': r[2] or 0, 'turnover_rate': r[3] or 0}
            for r in top_hot_result.fetchall()
        ]
        
        # TOP冷门库位 - 筛选指定库区（按库位分组去重）
        if zone_id:
            top_cold_result = await db.execute(
                select(
                    Location.full_code,
                    func.min(LocationHeatData.heat_value),
                    func.sum(LocationHeatData.pick_frequency),
                    func.avg(LocationHeatData.turnover_rate)
                ).join(Location, LocationHeatData.location_id == Location.id)
                .where(LocationHeatData.location_id.in_(location_ids_query))
                .group_by(Location.id, Location.full_code)
                .order_by(func.min(LocationHeatData.heat_value).asc())
                .limit(20)
            )
        else:
            top_cold_result = await db.execute(
                select(
                    Location.full_code,
                    func.min(LocationHeatData.heat_value),
                    func.sum(LocationHeatData.pick_frequency),
                    func.avg(LocationHeatData.turnover_rate)
                ).join(Location, LocationHeatData.location_id == Location.id)
                .group_by(Location.id, Location.full_code)
                .order_by(func.min(LocationHeatData.heat_value).asc())
                .limit(20)
            )
        data['top_cold'] = [
            {'full_code': r[0], 'heat_value': r[1] or 0, 'pick_frequency': r[2] or 0, 'turnover_rate': r[3] or 0}
            for r in top_cold_result.fetchall()
        ]
        
        # 巷道分析 - 筛选指定库区
        if zone_id:
            aisle_result = await db.execute(
                select(
                    Zone.code,
                    Aisle.code,
                    Aisle.name,
                    func.avg(LocationHeatData.heat_value),
                    func.sum(LocationHeatData.pick_frequency),
                    func.count(func.distinct(Location.id))  # 去重统计库位数
                ).join(Zone, Aisle.zone_id == Zone.id)
                .join(Shelf, Aisle.id == Shelf.aisle_id)
                .join(Location, Shelf.id == Location.shelf_id)
                .join(LocationHeatData, Location.id == LocationHeatData.location_id)
                .where(Aisle.zone_id == zone_id)
                .group_by(Zone.code, Aisle.id, Aisle.code, Aisle.name)
                .order_by(func.avg(LocationHeatData.heat_value).desc())
            )
        else:
            aisle_result = await db.execute(
                select(
                    Zone.code,
                    Aisle.code,
                    Aisle.name,
                    func.avg(LocationHeatData.heat_value),
                    func.sum(LocationHeatData.pick_frequency),
                    func.count(func.distinct(Location.id))  # 去重统计库位数
                ).join(Zone, Aisle.zone_id == Zone.id)
                .join(Shelf, Aisle.id == Shelf.aisle_id)
                .join(Location, Shelf.id == Location.shelf_id)
                .join(LocationHeatData, Location.id == LocationHeatData.location_id)
                .group_by(Zone.code, Aisle.id, Aisle.code, Aisle.name)
                .order_by(func.avg(LocationHeatData.heat_value).desc())
            )
        data['aisle_analysis'] = [
            {'zone_code': r[0], 'code': r[1], 'name': r[2], 'avg_heat': r[3] or 0, 'total_freq': r[4] or 0, 'loc_cnt': r[5]}
            for r in aisle_result.fetchall()
        ]
        
        # 货架分析（TOP 15 热门货架）- 筛选指定库区
        if zone_id:
            shelf_result = await db.execute(
                select(
                    Zone.code,
                    Aisle.code,
                    Shelf.code,
                    func.avg(LocationHeatData.heat_value),
                    func.sum(LocationHeatData.pick_frequency),
                    func.count(func.distinct(Location.id))  # 去重统计库位数
                ).join(Zone, Aisle.zone_id == Zone.id)
                .join(Shelf, Aisle.id == Shelf.aisle_id)
                .join(Location, Shelf.id == Location.shelf_id)
                .join(LocationHeatData, Location.id == LocationHeatData.location_id)
                .where(Aisle.zone_id == zone_id)
                .group_by(Zone.code, Aisle.code, Shelf.id, Shelf.code)
                .order_by(func.avg(LocationHeatData.heat_value).desc())
                .limit(15)
            )
        else:
            shelf_result = await db.execute(
                select(
                    Zone.code,
                    Aisle.code,
                    Shelf.code,
                    func.avg(LocationHeatData.heat_value),
                    func.sum(LocationHeatData.pick_frequency),
                    func.count(func.distinct(Location.id))  # 去重统计库位数
                ).join(Zone, Aisle.zone_id == Zone.id)
                .join(Shelf, Aisle.id == Shelf.aisle_id)
                .join(Location, Shelf.id == Location.shelf_id)
                .join(LocationHeatData, Location.id == LocationHeatData.location_id)
                .group_by(Zone.code, Aisle.code, Shelf.id, Shelf.code)
                .order_by(func.avg(LocationHeatData.heat_value).desc())
                .limit(15)
            )
        data['shelf_analysis'] = [
            {'zone_code': r[0], 'aisle_code': r[1], 'shelf_code': r[2], 'avg_heat': r[3] or 0, 'total_freq': r[4] or 0, 'loc_cnt': r[5]}
            for r in shelf_result.fetchall()
        ]
        
        return data


def generate_docx_report(data: dict, output_path: str) -> bool:
    """生成Word文档报告"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装，无法生成报告")
    
    doc = Document()
    
    # 设置文档默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    report_scope = data.get('report_scope', '全部')
    title = doc.add_heading('仓库库位热力图分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 分析范围
    scope_para = doc.add_paragraph()
    scope_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope_para.add_run(f'分析范围：{report_scope}').bold = True
    
    # 报告信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}').italic = True
    
    doc.add_paragraph()
    
    # ===== 第一章：数据概览 =====
    doc.add_heading('一、数据概览', level=1)
    
    # 基础统计表格
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    stats = [
        ('指标', '数值'),
        ('仓库总数', f"{data['warehouse_count']} 个"),
        ('库区总数', f"{data['zone_count']} 个"),
        ('巷道总数', f"{data['aisle_count']} 条"),
        ('货架总数', f"{data['shelf_count']} 个"),
        ('库位总数', f"{data['location_count']} 个"),
        ('有热力数据的库位', f"{data['active_location_count']} 个"),
        ('热力数据记录', f"{data['heat_data_count']} 条"),
    ]
    
    for i, (key, val) in enumerate(stats):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = str(val)
        if i == 0:
            for cell in table.rows[i].cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # ===== 第二章：热度分析 =====
    doc.add_heading('二、热度分析', level=1)
    
    heat_stats = data['heat_stats']
    freq_stats = data['freq_stats']
    
    doc.add_paragraph(f"热度值范围：{heat_stats['min']:.2f} ~ {heat_stats['max']:.2f}")
    doc.add_paragraph(f"平均热度：{heat_stats['avg']:.2f}")
    doc.add_paragraph(f"拣货频率范围：{freq_stats['min']} ~ {freq_stats['max']}")
    doc.add_paragraph(f"平均拣货频率：{freq_stats['avg']:.2f}")
    
    doc.add_heading('热度分布统计', level=2)
    
    # 热度分布表格
    dist_data = data['heat_distribution']
    total = sum(d['cnt'] for d in dist_data)
    
    table = doc.add_table(rows=len(dist_data) + 1, cols=3)
    table.style = 'Table Grid'
    
    headers = ['热度等级', '库位数量', '占比']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for i, d in enumerate(dist_data, 1):
        pct = d['cnt'] / total * 100 if total > 0 else 0
        table.rows[i].cells[0].text = d['heat_level']
        table.rows[i].cells[1].text = str(d['cnt'])
        table.rows[i].cells[2].text = f"{pct:.1f}%"
    
    doc.add_paragraph()
    
    # ===== 第三章：巷道热度分析 =====
    doc.add_heading('三、巷道热度分析', level=1)
    
    doc.add_paragraph('按巷道统计各区域的热度分布情况，热度越高表示该巷道的拣货频率越高。')
    doc.add_paragraph()
    
    aisle_data = data['aisle_analysis']
    
    if aisle_data:
        table = doc.add_table(rows=len(aisle_data) + 1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['巷道', '平均热度', '总拣货频率', '库位数']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for i, d in enumerate(aisle_data, 1):
            # 巷道名称加上库区前缀，格式：LP-01巷
            table.rows[i].cells[0].text = f"{d['zone_code']}-{d['code']}"
            table.rows[i].cells[1].text = f"{d['avg_heat']:.2f}"
            table.rows[i].cells[2].text = str(int(d['total_freq']))
            table.rows[i].cells[3].text = str(d['loc_cnt'])
    
    doc.add_paragraph()
    
    # ===== 第三章补充：热门货架分析 =====
    doc.add_heading('热门货架 TOP 15', level=2)
    
    shelf_data = data['shelf_analysis']
    
    if shelf_data:
        table = doc.add_table(rows=len(shelf_data) + 1, cols=5)
        table.style = 'Table Grid'
        
        headers = ['巷道', '货架', '平均热度', '总拣货频率', '库位数']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for i, d in enumerate(shelf_data, 1):
            # 巷道名称加上库区前缀，格式：LP-01巷
            table.rows[i].cells[0].text = f"{d['zone_code']}-{d['aisle_code']}"
            table.rows[i].cells[1].text = d['shelf_code']
            table.rows[i].cells[2].text = f"{d['avg_heat']:.2f}"
            table.rows[i].cells[3].text = str(int(d['total_freq']))
            table.rows[i].cells[4].text = str(d['loc_cnt'])
    
    doc.add_paragraph()
    
    # ===== 第四章：TOP库位 =====
    doc.add_heading('四、库位热度排名', level=1)
    
    doc.add_heading('TOP 10 热门库位', level=2)
    hot_data = data['top_hot'][:10]
    
    table = doc.add_table(rows=len(hot_data) + 1, cols=4)
    table.style = 'Table Grid'
    
    headers = ['库位编码', '热度值', '拣货频率', '周转率']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for i, d in enumerate(hot_data, 1):
        table.rows[i].cells[0].text = d['full_code']
        table.rows[i].cells[1].text = f"{d['heat_value']:.2f}"
        table.rows[i].cells[2].text = str(d['pick_frequency'])
        table.rows[i].cells[3].text = f"{d['turnover_rate']:.4f}"
    
    doc.add_paragraph()
    
    doc.add_heading('TOP 10 冷门库位', level=2)
    cold_data = data['top_cold'][:10]
    
    table = doc.add_table(rows=len(cold_data) + 1, cols=4)
    table.style = 'Table Grid'
    
    headers = ['库位编码', '热度值', '拣货频率', '周转率']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for i, d in enumerate(cold_data, 1):
        table.rows[i].cells[0].text = d['full_code']
        table.rows[i].cells[1].text = f"{d['heat_value']:.2f}"
        table.rows[i].cells[2].text = str(d['pick_frequency'])
        table.rows[i].cells[3].text = f"{d['turnover_rate']:.4f}"
    
    doc.add_paragraph()
    
    # ===== 第五章：优化建议 =====
    doc.add_heading('五、优化建议', level=1)
    
    suggestions = [
        '【商品ABC分类调整】',
        '  - A类商品（前20%高频）：靠近出库口的巷道',
        '  - B类商品（中间30%）：中等位置巷道',
        '  - C类商品（后50%低频）：远端巷道',
        '',
        '【热区分流】',
        '  - 将热度>400的库位中部分SKU迁移至较冷巷道',
        '  - 将冷门巷道迁入B类商品，提升利用率',
        '  - 避免单一巷道过热导致拣货拥堵',
        '',
        '【巷道均衡优化】',
        '  - 热门巷道（热度>300）：控制SKU数量，分散到相邻巷道',
        '  - 冷门巷道（热度<100）：迁入中频商品提升利用率',
        '  - 保持各巷道热度相对均衡，提高整体拣货效率',
        '',
        '【动态调整机制】',
        '  - 每月复盘热力图数据',
        '  - 根据销售季节性调整库位',
        '  - 建立库位热度监控预警'
    ]
    
    for s in suggestions:
        doc.add_paragraph(s)
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('—— 报告结束 ——').italic = True
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.add_run('生成工具：WMS仓库热力图系统').font.size = Pt(9)
    
    # 保存文档
    doc.save(output_path)
    return True


@router.get("/generate")
async def generate_report(zone_id: Optional[int] = None):
    """生成热力分析报告"""
    try:
        # 获取数据
        data = await fetch_report_data(zone_id)
        
        # 生成报告文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        reports_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'reports')
        os.makedirs(reports_dir, exist_ok=True)
        
        filename = f'heatmap_report_{timestamp}.docx'
        output_path = os.path.join(reports_dir, filename)
        
        generate_docx_report(data, output_path)
        
        return {
            "success": True,
            "filename": filename,
            "message": "报告生成成功"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"报告生成失败: {str(e)}")


@router.get("/download/{filename}")
async def download_report(filename: str):
    """下载报告文件"""
    reports_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'reports')
    file_path = os.path.join(reports_dir, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="报告文件不存在")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@router.get("/list")
async def list_reports():
    """获取报告列表"""
    reports_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'reports')
    
    if not os.path.exists(reports_dir):
        return {"reports": []}
    
    reports = []
    for filename in os.listdir(reports_dir):
        if filename.endswith('.docx'):
            file_path = os.path.join(reports_dir, filename)
            stat = os.stat(file_path)
            reports.append({
                "filename": filename,
                "size": stat.st_size,
                "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat()
            })
    
    # 按创建时间倒序
    reports.sort(key=lambda x: x['created_at'], reverse=True)
    
    return {"reports": reports}
